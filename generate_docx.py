"""
Generate README.docx from README.md
Creates a beautiful Word document with all tables and placeholders for diagrams.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Professional Color Palette (Navy Blue theme)
COLORS = {
    'primary': '1E3A5F',      # Deep Navy Blue
    'secondary': '2E5077',     # Medium Navy
    'accent': '4A90A4',        # Teal Blue
    'light': 'E8F4F8',         # Light Blue-Gray
    'dark': '0F1C2E',          # Dark Navy
    'success': '2D6A4F',       # Forest Green
    'text': '2C3E50',          # Dark Slate
    'white': 'FFFFFF',
    'gray': 'F8FAFC',          # Very Light Gray
    'border': 'CBD5E1',        # Border Gray
}

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, header_color='primary'):
    """Add a beautifully styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header row with gradient-like styling
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        para = header_cells[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(header_cells[i], COLORS[header_color])
    
    # Data rows with alternating colors
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        bg_color = COLORS['gray'] if row_idx % 2 == 0 else COLORS['white']
        
        for col_idx, cell_text in enumerate(row_data):
            row_cells[col_idx].text = cell_text
            para = row_cells[col_idx].paragraphs[0]
            if para.runs:
                para.runs[0].font.size = Pt(10)
                para.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            set_cell_shading(row_cells[col_idx], bg_color)
            
            # Bold first column
            if col_idx == 0 and para.runs:
                para.runs[0].bold = True
    
    return table

def add_diagram_placeholder(doc, title):
    """Add a professional placeholder for a Mermaid diagram."""
    # Create a bordered box effect
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Placeholder box
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(f"📊 {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(30, 58, 95)  # Navy
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p3.add_run("[ Insert rendered diagram image here ]")
    run2.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(74, 144, 164)  # Teal
    doc.add_paragraph()

def add_section_heading(doc, text, level=1):
    """Add a styled section heading."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(30, 58, 95)  # Navy
    return heading

def add_code_block(doc, code_text):
    """Add a styled code block."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(44, 62, 80)
    # Add light background via paragraph shading would require more complex XML
    return p

def create_readme_docx():
    doc = Document()
    
    # ================= TITLE SECTION =================
    # Main title
    title = doc.add_heading('', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('🔒 Local LLM Policy Gap Analyzer')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(30, 58, 95)
    run.bold = True
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Privacy-First Cybersecurity Policy Analysis')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(74, 144, 164)
    run.italic = True
    
    # Tagline
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run('Against NIST CSF Standards')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(46, 80, 119)
    
    # Badge-like info
    badges = doc.add_paragraph()
    badges.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_text = '🐍 Python 3.8+  •  🦙 Ollama  •  🛡️ NIST CSF 2024  •  📡 100% Offline'
    run = badges.add_run(badge_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(45, 106, 79)
    run.bold = True
    
    doc.add_paragraph()
    
    # ================= INTRODUCTION =================
    add_section_heading(doc, '📖 Introduction', 1)
    
    intro1 = doc.add_paragraph()
    run = intro1.add_run(
        'Organizations struggle to maintain cybersecurity policies that align with industry standards. '
        'This tool provides '
    )
    run.font.size = Pt(11)
    run2 = intro1.add_run('automated gap analysis')
    run2.bold = True
    run2.font.size = Pt(11)
    run3 = intro1.add_run(
        ' by comparing your policies against the NIST Cybersecurity Framework '
        'using a local LLM (Gemma3 via Ollama).'
    )
    run3.font.size = Pt(11)
    
    intro2 = doc.add_paragraph()
    run = intro2.add_run('Every operation runs entirely on your machine—')
    run.font.size = Pt(11)
    run2 = intro2.add_run('no cloud APIs, no data collection, complete privacy.')
    run2.bold = True
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(45, 106, 79)
    
    doc.add_paragraph()
    
    # ================= TABLE OF CONTENTS =================
    add_section_heading(doc, '📋 Table of Contents', 1)
    toc_data = [
        ('Features', 'Core capabilities'),
        ('Tech Stack', 'Technologies and requirements'),
        ('Architecture', 'Visual system overview'),
        ('Project Structure', 'File organization'),
        ('Quick Start', 'Get running in 5 minutes'),
        ('Developer Guide', 'Contributing code'),
        ('Known Issues', 'Current limitations'),
    ]
    add_styled_table(doc, ['Section', 'Description'], toc_data, 'primary')
    doc.add_paragraph()
    
    # ================= FEATURES =================
    add_section_heading(doc, '⭐ Features', 1)
    features_data = [
        ('🔍 Gap Analysis', 'Identifies policy weaknesses against NIST CSF standards'),
        ('📝 Policy Revision', 'Auto-generates improved policy versions addressing gaps'),
        ('🗺️ Implementation Roadmap', 'Phased improvement plans (0-3, 3-6, 6-12 months)'),
        ('📊 Executive Summary', 'Leadership-ready overview of findings'),
        ('📁 Multi-Format Input', 'Supports .txt, .pdf, and .docx policies'),
        ('📄 PDF Output', 'Professional formatted reports using ReportLab'),
        ('⚡ Batch Processing', 'Analyze multiple policies in one run'),
        ('🔒 100% Offline', 'Zero network calls after initial setup'),
    ]
    add_styled_table(doc, ['Feature', 'Description'], features_data, 'success')
    doc.add_paragraph()
    
    # ================= TECH STACK =================
    add_section_heading(doc, '🛠️ Tech Stack & Prerequisites', 1)
    
    add_section_heading(doc, 'Technology Stack', 2)
    add_diagram_placeholder(doc, 'Technology Stack Architecture')
    
    add_section_heading(doc, 'System Requirements', 2)
    requirements_data = [
        ('💻 CPU', 'Intel i5 / AMD Ryzen 5', 'Intel i7 / AMD Ryzen 7'),
        ('🧠 RAM', '8 GB', '16 GB'),
        ('💾 Storage', '10 GB', '20 GB'),
        ('🖥️ OS', 'Windows 10 / Linux / macOS', 'Windows 11 / Ubuntu 22.04'),
    ]
    add_styled_table(doc, ['Component', 'Minimum', 'Recommended'], requirements_data, 'secondary')
    doc.add_paragraph()
    
    add_section_heading(doc, 'Dependencies', 2)
    deps_data = [
        ('PyPDF2', '>= 3.0', 'PDF text extraction'),
        ('python-docx', '>= 0.8', 'Word document parsing'),
        ('reportlab', '>= 4.0', 'PDF report generation'),
        ('ollama', '(runtime)', 'Local LLM execution'),
    ]
    add_styled_table(doc, ['Package', 'Version', 'Purpose'], deps_data, 'accent')
    doc.add_paragraph()
    
    # ================= ARCHITECTURE =================
    add_section_heading(doc, '🏗️ Architecture', 1)
    
    add_section_heading(doc, 'System Overview', 2)
    add_diagram_placeholder(doc, 'System Architecture Flowchart')
    
    add_section_heading(doc, 'Data Flow', 2)
    add_diagram_placeholder(doc, 'Data Flow Pipeline')
    
    add_section_heading(doc, 'Processing Sequence', 2)
    add_diagram_placeholder(doc, 'UML Sequence Diagram')
    
    add_section_heading(doc, 'NIST CSF Coverage', 2)
    add_diagram_placeholder(doc, 'NIST Framework Map')
    
    # ================= PROJECT STRUCTURE =================
    add_section_heading(doc, '📂 Project Structure', 1)
    
    structure_text = """Local-LLM/
├── src/
│   ├── main.py              # CLI & orchestrator
│   ├── gap_analyzer.py      # NIST comparison & LLM
│   ├── policy_reviser.py    # Policy improvements
│   ├── roadmap_generator.py # Roadmap creation
│   ├── pdf_generator.py     # PDF formatting
│   └── utils.py             # File utilities
├── data/
│   ├── reference/           # NIST CSF files
│   └── test_policies/       # Sample policies
├── output/                  # Generated reports
└── requirements.txt         # Dependencies"""
    
    add_code_block(doc, structure_text)
    doc.add_paragraph()
    
    add_section_heading(doc, 'Module Overview', 2)
    modules_data = [
        ('main.py', '216 lines', 'CLI parsing, workflow orchestration'),
        ('gap_analyzer.py', '131 lines', 'NIST loading, policy comparison'),
        ('policy_reviser.py', '65 lines', 'Policy improvement generation'),
        ('roadmap_generator.py', '112 lines', 'Phased roadmap creation'),
        ('pdf_generator.py', '167 lines', 'PDF formatting with ReportLab'),
        ('utils.py', '78 lines', 'Document reading utilities'),
    ]
    add_styled_table(doc, ['Module', 'Size', 'Responsibility'], modules_data, 'primary')
    doc.add_paragraph()
    
    # ================= QUICK START =================
    add_section_heading(doc, '🚀 Quick Start', 1)
    
    add_section_heading(doc, 'Step 1: Clone & Setup', 2)
    install_code = """git clone https://github.com/HACK-IITK-2025-C3iHub/Local-LLM.git
cd "Local LLM"
python -m venv venv
venv\\Scripts\\activate
pip install -r requirements.txt"""
    add_code_block(doc, install_code)
    doc.add_paragraph()
    
    add_section_heading(doc, 'Step 2: Install Ollama', 2)
    ollama_data = [
        ('1️⃣ Install', 'Download from ollama.ai', 'One-time'),
        ('2️⃣ Pull Model', 'ollama pull gemma3:latest', 'Needs internet'),
        ('3️⃣ Verify', 'ollama list', 'Should show gemma3'),
    ]
    add_styled_table(doc, ['Step', 'Command', 'Notes'], ollama_data, 'success')
    doc.add_paragraph()
    
    add_section_heading(doc, 'Step 3: Run Analysis', 2)
    run_code = """# Single policy
python src/main.py --policy data/test_policies/isms_policy.txt

# Batch processing
python src/main.py --batch data/test_policies/"""
    add_code_block(doc, run_code)
    doc.add_paragraph()
    
    add_section_heading(doc, 'Output Reports', 2)
    reports_data = [
        ('📊 Gap Analysis', 'TXT + PDF', 'Identified weaknesses'),
        ('📝 Revised Policy', 'TXT + PDF', 'Improved version'),
        ('🗺️ Roadmap', 'TXT + PDF', 'Implementation plan'),
        ('📋 Executive Summary', 'TXT + PDF', 'Leadership overview'),
        ('📚 Comprehensive', 'TXT + PDF', 'All combined'),
    ]
    add_styled_table(doc, ['Report', 'Format', 'Description'], reports_data, 'accent')
    doc.add_paragraph()
    
    add_section_heading(doc, 'Processing Time', 2)
    time_data = [
        ('Gap Analysis', '1-2 min'),
        ('Policy Revision', '2-3 min'),
        ('Roadmap Generation', '1-2 min'),
        ('Executive Summary', '30-60 sec'),
        ('TOTAL', '~5-8 min'),
    ]
    add_styled_table(doc, ['Stage', 'Duration'], time_data, 'primary')
    doc.add_paragraph()
    
    # ================= DEVELOPER GUIDE =================
    add_section_heading(doc, '👨‍💻 Developer Guide', 1)
    
    add_section_heading(doc, 'Code Architecture', 2)
    add_diagram_placeholder(doc, 'Code Structure Flowchart')
    
    add_section_heading(doc, 'Security Limits', 2)
    limits_data = [
        ('LLM_TIMEOUT', '600 seconds'),
        ('MAX_PROMPT_SIZE', '100 KB'),
        ('MAX_POLICY_SIZE', '50 KB'),
        ('MAX_FILE_SIZE', '50 MB'),
    ]
    add_styled_table(doc, ['Constant', 'Value'], limits_data, 'secondary')
    doc.add_paragraph()
    
    # ================= KNOWN ISSUES =================
    add_section_heading(doc, '⚠️ Known Issues & Limitations', 1)
    
    limitations_data = [
        ('LLM Accuracy', 'Output depends on model', 'Review manually'),
        ('Processing Time', '5-8 min per policy', 'Batch overnight'),
        ('RAM Usage', '~6GB for Gemma3', 'Close other apps'),
        ('Language', 'English only', 'Translation planned'),
    ]
    add_styled_table(doc, ['Issue', 'Description', 'Workaround'], limitations_data, 'primary')
    doc.add_paragraph()
    
    # ================= FOOTER =================
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Divider line
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run('━' * 40)
    run.font.color.rgb = RGBColor(203, 213, 225)
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Made With ❤️ by T-reXploit')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(30, 58, 95)
    
    sub_footer = doc.add_paragraph()
    sub_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_footer.add_run('🔒 Privacy-First  •  📡 Offline-Ready  •  💻 Open Source')
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(74, 144, 164)
    
    # GitHub
    github = doc.add_paragraph()
    github.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = github.add_run('github.com/HACK-IITK-2025-C3iHub/Local-LLM')
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(107, 114, 128)
    
    # Save
    output_path = 'README.docx'
    doc.save(output_path)
    
    print("=" * 50)
    print("✅ Created README.docx with professional styling!")
    print("=" * 50)
    print("\n🎨 Color Theme: Navy Blue Professional")
    print("📊 Tables: 10 styled tables with alternating rows")
    print("📌 Placeholders: 6 diagram placeholders ready for images")
    print("\n📝 Next Steps:")
    print("   1. Open README.docx")
    print("   2. Screenshot diagrams from GitHub")
    print("   3. Replace placeholders with images")
    
    return output_path

if __name__ == "__main__":
    create_readme_docx()
